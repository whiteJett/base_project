from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import docx
from app.config import settings

def load_pdf(path: Path) -> List[Document]:
    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(path), "page": i+1}
            ))
    return docs

def load_docx(path: Path) -> List[Document]:
    d = docx.Document(str(path))
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": str(path)})] if text else []

def load_docs(dir_path: str) -> List[Document]:
    p = Path(dir_path)
    docs: List[Document] = []
    for f in p.rglob("*"):
        if f.suffix.lower() == ".pdf":
            docs.extend(load_pdf(f))
        elif f.suffix.lower() in [".docx", ".doc"]:
            docs.extend(load_docx(f))
        elif f.suffix.lower() in [".md", ".txt"]:
            docs.append(Document(page_content=f.read_text(encoding="utf-8"),
                                 metadata={"source": str(f)}))
    return docs

def split_docs(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap
    )
    return splitter.split_documents(docs)

if __name__ == "__main__":
    docs=split_docs(load_docs("/home/zero/PycharmProjects/PythonProject1/date/docs"))
    for _ in docs:
        print(_)
